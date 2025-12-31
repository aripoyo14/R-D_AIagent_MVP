"""
サイドバーコンポーネント
設定とAPIキー確認、面談情報入力を表示する
"""

import streamlit as st
import os
from services.ai_review import review_interview_content
from typing import Dict, Tuple, Optional
import io
import docx
import pypdf

# 事業部のリスト
DEPARTMENTS = [
    "製品開発部",
    "研究開発部",
    "技術営業部",
    "材料技術部"
]


def check_api_keys() -> bool:
    """APIキーの設定状況を確認"""
    try:
        supabase_url = os.getenv("SUPABASE_URL")
        supabase_key = os.getenv("SUPABASE_KEY")
        openai_api_key = os.getenv("OPENAI_API_KEY")
        
        has_supabase = supabase_url is not None and supabase_url != "" and supabase_key is not None and supabase_key != ""
        has_openai = openai_api_key is not None and openai_api_key != ""
        return has_supabase and has_openai
    except:
        return False


def render_sidebar(review_container: Optional[st.delta_generator.DeltaGenerator] = None) -> Tuple[str, bool, Dict]:
    """
    サイドバーを表示する
    
    Returns:
        tuple: (選択された事業部名, APIキー設定状況, フォームデータ)
    """
    # Initialize variables
    selected_department = DEPARTMENTS[0]
    api_keys_ok = False
    form_data = {}
    model_name = "gemini-2.5-flash-lite" # Default

    # ロゴを中央揃えで表示
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.image("images/AgentX_logo.png", use_container_width=True)

    # タブを作成
    tab1, tab2 = st.tabs(["📝 面談情報入力", "⚙️ 設定"])
    
    # タブ2: 設定 (先にレンダリングしてmodel_nameを取得)
    with tab2:
        # 事業部選択
        selected_department = st.selectbox(
            "事業部を選択",
            DEPARTMENTS,
            index=0
        )
        
        st.divider()

        # AIモデル選択
        model_name = st.selectbox(
            "AIモデルを選択",
            ["gemini-2.5-pro", "gemini-2.5-flash", "gemini-2.5-flash-lite"],
            index=2  # Default to gemini-2.5-flash
        )
        
        st.divider()
        
        # APIキー設定状況
        api_keys_ok = check_api_keys()

    # タブ1: 面談情報入力 (取得したmodel_nameを使用)
    with tab1:
        form_data = render_interview_form(review_container, model_name=model_name)
    
    return selected_department, api_keys_ok, form_data, model_name


def render_interview_form(review_container: Optional[st.delta_generator.DeltaGenerator] = None, model_name: str = "gemini-2.5-flash-lite") -> Dict:
    """
    面談情報入力フォームを表示する
    
    Returns:
        Dict: フォームデータ（company_name, contact_info, interview_memo, submitted）
    """
    # フォーム定義を削除し、各要素を直接配置することでファイルアップロードの即時反映を実現
    # with st.form("interview_form", clear_on_submit=False):
    
    company_name = st.text_input(
        "企業名 (Company Name)",
        value=st.session_state.form_data.get("company_name", ""),
        placeholder="例: サンプル自動車"
    )
    
    contact_info = st.text_input(
        "相手方 部署・役職",
        value=st.session_state.form_data.get("contact_info", ""),
        placeholder="例: ボディ設計部 課長"
    )
    
    # 面談メモはファイルアップロードからのみ取得
    interview_memo = st.session_state.form_data.get("interview_memo", "")

    if not interview_memo:
        uploaded_file = st.file_uploader(
            "ファイルから読み込む (docx, txt, pdf)",
            type=["docx", "txt", "pdf"],
            key="interview_file_uploader_sidebar"
        )

        if uploaded_file is not None:
            try:
                text = ""
                if uploaded_file.type == "text/plain":
                    text = uploaded_file.getvalue().decode("utf-8")
                elif uploaded_file.type == "application/pdf":
                    pdf_reader = pypdf.PdfReader(uploaded_file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    doc = docx.Document(uploaded_file)
                    for para in doc.paragraphs:
                        text += para.text + "\n"
                
                if text:
                    st.session_state.form_data["interview_memo"] = text
                    st.rerun()
            except Exception as e:
                st.error(f"ファイルの読み込みに失敗しました: {e}")
        
        st.info("👆 ファイルをアップロードしてください")
    else:
        st.success(f"✅ 面談メモを読み込みました ({len(interview_memo)}文字)")
        with st.expander("読み込んだ内容を確認"):
            st.text(interview_memo)
        
        if st.button("ファイルを削除 (Clear)"):
            st.session_state.form_data["interview_memo"] = ""
            st.rerun()
    
    submitted = st.button("AIレビュー実行", type="primary", use_container_width=True)
    
    if submitted:
        if not interview_memo.strip():
            st.error("⚠️ 面談メモを入力してください")
        else:
            # フォームデータを保存
            st.session_state.form_data = {
                "company_name": company_name,
                "contact_info": contact_info,
                "interview_memo": interview_memo
            }
            
            # AIレビューを実行
            spinner_target = review_container or st
            with spinner_target:
                with st.spinner("🤖 AIが内容をレビュー中..."):
                    # 再実行のためにフラグをリセット
                    st.session_state.show_idea_report = False
                    st.session_state.is_agent_running = False
                    
                    review_result = review_interview_content(interview_memo, model_name=model_name)
                    st.session_state.review_result = review_result
    
    # デモ用面談録の読み込みとAIレビュー実行ボタン
    st.markdown("---")
    st.markdown("### 🎬 デモ用")
    
    demo_file_path = "demo_document.docx"
    if os.path.exists(demo_file_path):
        if st.button("📄 デモ用面談録を読み込んでAIレビュー実行", type="secondary", use_container_width=True):
            try:
                # デモ用ファイルを読み込む
                doc = docx.Document(demo_file_path)
                text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                
                if text:
                    # セッションステートに設定
                    st.session_state.form_data = {
                        "company_name": "サンプル株式会社",
                        "contact_info": "ロボティクス開発本部 ハードウェア設計部 佐藤 チーフアーキテクト、ジェニファー・ウー 製造技術マネージャー",
                        "interview_memo": text
                    }
                    
                    # AIレビューを自動実行
                    spinner_target = review_container or st
                    with spinner_target:
                        with st.spinner("🤖 AIが内容をレビュー中..."):
                            # 再実行のためにフラグをリセット
                            st.session_state.show_idea_report = False
                            st.session_state.is_agent_running = False
                            
                            review_result = review_interview_content(text, model_name=model_name)
                            st.session_state.review_result = review_result
                    
                    st.success("✅ デモ用面談録を読み込み、AIレビューを実行しました")
                    st.rerun()
            except Exception as e:
                st.error(f"デモ用ファイルの読み込みに失敗しました: {e}")
    else:
        st.info("ℹ️ デモ用ファイルが見つかりません")
    
    return {
        "company_name": company_name,
        "contact_info": contact_info,
        "interview_memo": interview_memo,
        "submitted": submitted
    }
