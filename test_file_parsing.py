import docx
import os

def create_test_files():
    # Create TXT
    with open("test.txt", "w", encoding="utf-8") as f:
        f.write("This is a test text file.")

    # Create DOCX
    doc = docx.Document()
    doc.add_paragraph("This is a test docx file.")
    doc.save("test.docx")

    print("Created test.txt and test.docx")

def test_parsing():
    # Test TXT
    with open("test.txt", "rb") as f:
        content = f.read().decode("utf-8")
        print(f"TXT Content: {content}")
        assert "This is a test text file." in content

    # Test DOCX
    doc = docx.Document("test.docx")
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    print(f"DOCX Content: {text}")
    assert "This is a test docx file." in text

    print("Verification successful!")

if __name__ == "__main__":
    create_test_files()
    test_parsing()
    # Cleanup
    if os.path.exists("test.txt"):
        os.remove("test.txt")
    if os.path.exists("test.docx"):
        os.remove("test.docx")
